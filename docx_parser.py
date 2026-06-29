"""
Parse Verbatim-style and generic debate .docx files into structured DebateCard objects.

Verbatim (and CardMirror) documents use named paragraph styles:
  Pocket > Hat > Block > Tag / Cite / CardText / Analytics

Generic debate docs use formatting heuristics:
  Short bold/underlined paragraph → tag
  Italic or Author-Year paragraph → cite
  Following body text → card evidence
"""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn

from models import CardType, DebateCard, ParsedDocument, Side

# ---------------------------------------------------------------------------
# Style name constants
# ---------------------------------------------------------------------------

_POCKET_STYLES = {"pocket", "pocket label"}
_HAT_STYLES = {"hat", "hat label"}
_BLOCK_STYLES = {"block", "block label", "blocklabel", "block title"}
_TAG_STYLES = {"tag", "verbatim tag", "cardtag"}
_CITE_STYLES = {"cite", "verbatim cite", "cardcite", "citation"}
_CARD_STYLES = {
    "verbatim", "cardtext", "card text", "card body",
    "underline", "hl", "highlighted", "plain text",
    "body text", "normal",  # fallbacks
}
_ANALYTIC_STYLES = {"analytics", "verbatim analytics", "analytic", "analytic text"}
_HEADING_PREFIXES = ("heading",)

# Regex for detecting cites: "Smith 2023," or "Smith '23"
_CITE_AUTHOR_YEAR = re.compile(
    r"^[A-Z][a-zA-Zé\-]+\s+(?:\d{4}|'\d{2})[,\s—]", re.MULTILINE
)
# Publication signals found in cite paragraphs
_CITE_SIGNALS = frozenset([
    "journal", "professor", "phd", "university", "review", "policy",
    "foreign", "national", "institute", "press", "et al", "http", "www",
    "doi", "quarterly", "times", "post", "report", "fellow", "director",
    "chair", "senior", "assistant", "associate", "visiting", "research",
])


def _norm_style(para) -> str:
    """Normalized paragraph style name, lower-cased."""
    if para.style and para.style.name:
        return para.style.name.lower().strip()
    return ""


def _is_heading(para) -> Optional[int]:
    """Return heading level 1–6 or None."""
    sn = _norm_style(para)
    for level in range(1, 7):
        if sn == f"heading {level}":
            return level
    return None


def _para_is_empty(para) -> bool:
    return not para.text.strip()


def _run_formatting(para) -> dict:
    """Aggregate run-level formatting across a paragraph."""
    runs = [r for r in para.runs if r.text.strip()]
    if not runs:
        return {"all_bold": False, "all_italic": False, "has_underline": False,
                "has_bold": False, "has_italic": False}
    return {
        "all_bold": all(bool(r.bold) for r in runs),
        "all_italic": all(bool(r.italic) for r in runs),
        "has_underline": any(bool(r.underline) for r in runs),
        "has_bold": any(bool(r.bold) for r in runs),
        "has_italic": any(bool(r.italic) for r in runs),
    }


def _looks_like_tag(text: str, fmt: dict) -> bool:
    """Heuristic for non-Verbatim tag detection."""
    if not text or len(text) > 300:
        return False
    if fmt["all_bold"] and not fmt["all_italic"] and len(text) < 250:
        return True
    if text.upper() == text and len(text) < 150 and len(text) > 5:
        return True
    if fmt["has_bold"] and fmt["has_underline"] and len(text) < 300:
        return True
    return False


def _looks_like_cite(text: str) -> bool:
    """Heuristic for non-Verbatim cite detection."""
    if not text or len(text) > 500:
        return False
    if _CITE_AUTHOR_YEAR.match(text[:80]):
        return True
    tl = text.lower()
    signal_count = sum(1 for s in _CITE_SIGNALS if s in tl)
    return signal_count >= 2 and len(text) < 400


def _card_id(file_path: str, index: int) -> str:
    h = hashlib.md5(f"{file_path}:{index}".encode()).hexdigest()[:10]
    return f"card_{h}"


def _is_verbatim_doc(doc: Document) -> bool:
    """Return True if document uses Verbatim paragraph styles."""
    verbatim_markers = _POCKET_STYLES | _HAT_STYLES | _BLOCK_STYLES | _TAG_STYLES | _CITE_STYLES | _ANALYTIC_STYLES
    checked = 0
    for para in doc.paragraphs:
        sn = _norm_style(para)
        if any(vs == sn or sn.startswith(vs) for vs in verbatim_markers):
            return True
        checked += 1
        if checked > 200:
            break
    return False


# ---------------------------------------------------------------------------
# Parser state machine
# ---------------------------------------------------------------------------

@dataclass
class _ParserState:
    file_path: str
    file_name: str
    is_verbatim: bool

    # Current hierarchy
    heading: str = ""
    heading_level: int = 0
    pocket: str = ""
    hat: str = ""
    block: str = ""

    # Current card being assembled
    tag: str = ""
    cite: str = ""
    card_text_parts: List[str] = field(default_factory=list)
    card_start_idx: int = 0
    in_card: bool = False

    # Accumulator
    cards: List[DebateCard] = field(default_factory=list)
    analytics: List[dict] = field(default_factory=list)
    headings: List[str] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)
    para_index: int = 0

    def _flush_card(self):
        """Finalize the current card and add to the list."""
        if not self.tag and not self.card_text_parts:
            return
        card_text = "\n".join(self.card_text_parts).strip()
        tag = self.tag.strip()
        cite = self.cite.strip()
        if not tag and not card_text:
            return
        # If no tag, use first line of card text
        if not tag and card_text:
            tag = card_text[:80]

        full = " | ".join(filter(None, [tag, cite, card_text[:200]]))
        card = DebateCard(
            id=_card_id(self.file_path, self.card_start_idx),
            tag=tag,
            cite=cite,
            card_text=card_text,
            full_text=f"{tag} {cite} {card_text}".strip(),
            card_type=CardType.EVIDENCE,
            pocket=self.pocket,
            hat=self.hat,
            block=self.block,
            source_file=self.file_name,
            source_file_path=self.file_path,
            paragraph_index=self.card_start_idx,
        )
        self.cards.append(card)
        self.tag = ""
        self.cite = ""
        self.card_text_parts = []
        self.in_card = False

    def start_card(self, tag: str, idx: int):
        self._flush_card()
        self.tag = tag
        self.cite = ""
        self.card_text_parts = []
        self.card_start_idx = idx
        self.in_card = True

    def set_cite(self, cite: str):
        self.cite = cite

    def add_card_text(self, text: str):
        if text:
            self.card_text_parts.append(text)
        self.in_card = True

    def add_analytic(self, text: str, idx: int):
        if text:
            self.analytics.append({
                "text": text,
                "pocket": self.pocket,
                "hat": self.hat,
                "block": self.block,
                "source_file": self.file_name,
                "paragraph_index": idx,
            })

    def set_pocket(self, text: str):
        self._flush_card()
        self.pocket = text
        self.hat = ""
        self.block = ""

    def set_hat(self, text: str):
        self._flush_card()
        self.hat = text
        self.block = ""

    def set_block(self, text: str):
        self._flush_card()
        self.block = text

    def set_heading(self, text: str, level: int):
        if level <= 2:
            self._flush_card()
            self.pocket = ""
            self.hat = ""
            self.block = ""
        self.heading = text
        self.heading_level = level
        if text not in self.headings:
            self.headings.append(text)

    def finish(self):
        self._flush_card()


def _process_verbatim(para, state: _ParserState):
    """Process a paragraph from a Verbatim-formatted document."""
    sn = _norm_style(para)
    text = para.text.strip()
    idx = state.para_index

    if _is_heading(para) is not None:
        state.set_heading(text, _is_heading(para))
    elif any(sn == s or sn.startswith(s) for s in _POCKET_STYLES):
        state.set_pocket(text)
    elif any(sn == s or sn.startswith(s) for s in _HAT_STYLES):
        state.set_hat(text)
    elif any(sn == s or sn.startswith(s) for s in _BLOCK_STYLES):
        state.set_block(text)
    elif any(sn == s or sn.startswith(s) for s in _TAG_STYLES):
        if text:
            state.start_card(text, idx)
    elif any(sn == s or sn.startswith(s) for s in _CITE_STYLES):
        if text:
            state.set_cite(text)
    elif any(sn == s or sn.startswith(s) for s in _ANALYTIC_STYLES):
        if text:
            state.add_analytic(text, idx)
    elif any(sn == s or sn.startswith(s) for s in _CARD_STYLES):
        if text:
            state.add_card_text(text)
    else:
        # Unknown style — add to card text if we're in a card
        if text and state.in_card:
            state.add_card_text(text)


def _process_generic(para, state: _ParserState):
    """Process a paragraph from a non-Verbatim document using heuristics."""
    text = para.text.strip()
    idx = state.para_index

    if not text:
        return

    heading_level = _is_heading(para)
    if heading_level is not None:
        state.set_heading(text, heading_level)
        return

    fmt = _run_formatting(para)

    if _looks_like_tag(text, fmt):
        state.start_card(text, idx)
    elif _looks_like_cite(text) and state.in_card and not state.cite:
        state.set_cite(text)
    elif state.in_card:
        state.add_card_text(text)
    else:
        # Before first tag — treat as analytic/header text
        if len(text) > 20:
            state.add_analytic(text, idx)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_debate_docx(file_path: str) -> ParsedDocument:
    """
    Parse a debate .docx file into structured DebateCard objects.

    Returns a ParsedDocument with all detected cards, analytics, and headings.
    """
    path = Path(file_path)
    if not path.exists():
        return ParsedDocument(
            file_path=file_path,
            file_name=path.name,
            parse_errors=[f"File not found: {file_path}"],
        )

    try:
        doc = Document(file_path)
    except Exception as e:
        return ParsedDocument(
            file_path=file_path,
            file_name=path.name,
            parse_errors=[f"Failed to open DOCX: {e}"],
        )

    is_verbatim = _is_verbatim_doc(doc)
    state = _ParserState(
        file_path=file_path,
        file_name=path.name,
        is_verbatim=is_verbatim,
    )

    for i, para in enumerate(doc.paragraphs):
        state.para_index = i
        if _para_is_empty(para):
            continue
        try:
            if is_verbatim:
                _process_verbatim(para, state)
            else:
                _process_generic(para, state)
        except Exception as e:
            state.errors.append(f"Para {i}: {e}")

    # Also process paragraphs inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    state.para_index += 1
                    if _para_is_empty(para):
                        continue
                    try:
                        if is_verbatim:
                            _process_verbatim(para, state)
                        else:
                            _process_generic(para, state)
                    except Exception as e:
                        state.errors.append(f"Table para: {e}")

    state.finish()

    return ParsedDocument(
        file_path=file_path,
        file_name=path.name,
        total_cards=len(state.cards),
        total_analytics=len(state.analytics),
        headings=state.headings,
        cards=state.cards,
        analytics=state.analytics,
        is_verbatim_format=is_verbatim,
        metadata={
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        },
        parse_errors=state.errors,
    )
