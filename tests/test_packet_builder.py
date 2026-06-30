import sys
from pathlib import Path
import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from models import DebateCard, Side
from packet_builder import build_card_packet
from docx_parser import parse_debate_docx


def _card(id, tag, cite="", card_text="", pocket="Deterrence", block="AT: No War"):
    return DebateCard(
        id=id,
        tag=tag,
        cite=cite,
        card_text=card_text,
        full_text=f"{tag} {cite} {card_text}".strip(),
        pocket=pocket,
        block=block,
    )


class TestBuildPacket:
    def test_creates_file(self, tmp_path):
        result = build_card_packet(
            [_card("c1", "Tag one", card_text="Evidence text")],
            output_path=tmp_path / "out.docx",
        )
        assert result["success"] is True
        assert Path(result["path"]).exists()

    def test_card_count_in_result(self, tmp_path):
        cards = [
            _card("c1", "Tag one", card_text="Evidence one"),
            _card("c2", "Tag two", card_text="Evidence two"),
        ]
        result = build_card_packet(cards, output_path=tmp_path / "out.docx")
        assert result["card_count"] == 2

    def test_empty_cards_returns_error(self, tmp_path):
        result = build_card_packet([], output_path=tmp_path / "out.docx")
        assert result["success"] is False
        assert "error" in result

    def test_group_by_pocket(self, tmp_path):
        cards = [
            _card("c1", "Tag one", pocket="Aff Case", block=""),
            _card("c2", "Tag two", pocket="Neg Case", block=""),
            _card("c3", "Tag three", pocket="Aff Case", block=""),
        ]
        result = build_card_packet(
            cards, output_path=tmp_path / "out.docx", group_by="pocket"
        )
        assert result["success"] is True
        assert result["group_count"] == 2

    def test_group_by_block(self, tmp_path):
        cards = [
            _card("c1", "Tag one", block="Block A"),
            _card("c2", "Tag two", block="Block B"),
        ]
        result = build_card_packet(
            cards, output_path=tmp_path / "out.docx", group_by="block"
        )
        assert result["success"] is True
        assert result["group_count"] == 2

    def test_group_by_none(self, tmp_path):
        cards = [_card("c1", "Tag one"), _card("c2", "Tag two")]
        result = build_card_packet(
            cards, output_path=tmp_path / "out.docx", group_by="none"
        )
        assert result["success"] is True
        assert result["group_count"] == 1

    def test_group_by_source_file(self, tmp_path):
        cards = [
            DebateCard(id="c1", tag="Tag one", full_text="Tag one", source_file="file_a.docx"),
            DebateCard(id="c2", tag="Tag two", full_text="Tag two", source_file="file_b.docx"),
        ]
        result = build_card_packet(
            cards, output_path=tmp_path / "out.docx", group_by="source_file"
        )
        assert result["success"] is True
        assert result["group_count"] == 2

    def test_auto_output_path(self):
        cards = [_card("c1", "Tag", card_text="text")]
        result = build_card_packet(cards)
        assert result["success"] is True
        path = Path(result["path"])
        assert path.exists()
        path.unlink()

    def test_custom_title_appears_in_doc(self, tmp_path):
        cards = [_card("c1", "Tag", card_text="text")]
        result = build_card_packet(
            cards, output_path=tmp_path / "out.docx", title="My Custom Packet"
        )
        assert result["success"] is True
        doc = Document(result["path"])
        text = " ".join(p.text for p in doc.paragraphs)
        assert "My Custom Packet" in text

    def test_round_trip_text_preserved(self, tmp_path):
        cards = [
            _card("c1", "Deterrence prevents war", cite="Smith 2023",
                  card_text="Nuclear deterrence has maintained stability."),
        ]
        out = tmp_path / "packet.docx"
        build_card_packet(cards, output_path=out, include_source_metadata=False)
        parsed = parse_debate_docx(str(out))
        all_text = " ".join(c.full_text for c in parsed.cards)
        assert "deterrence" in all_text.lower()

    def test_round_trip_at_least_one_card(self, tmp_path):
        cards = [
            _card("c1", "Deterrence prevents war", card_text="Evidence one."),
            _card("c2", "Credibility is key", card_text="Evidence two."),
        ]
        out = tmp_path / "packet.docx"
        build_card_packet(cards, output_path=out, include_source_metadata=False)
        parsed = parse_debate_docx(str(out))
        assert parsed.total_cards >= 1

    def test_source_metadata_line_appears(self, tmp_path):
        cards = [
            DebateCard(
                id="c1", tag="Tag", full_text="Tag",
                team="Harvard KP", tournament="NDT 2024", source_file="test.docx",
            )
        ]
        result = build_card_packet(
            cards, output_path=tmp_path / "out.docx", include_source_metadata=True
        )
        doc = Document(result["path"])
        text = " ".join(p.text for p in doc.paragraphs)
        assert "Harvard KP" in text
