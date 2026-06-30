"""Standalone script to regenerate test fixture .docx files.

Run from repo root:
    python tests/fixtures/make_fixtures.py
"""
from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

FIXTURES = Path(__file__).parent


def _add_style(doc, name):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def make_verbatim(path: Path):
    """Verbatim-styled doc: Pocket/Hat/Block/Tag/Cite/Verbatim paragraph styles, 2 cards."""
    doc = Document()
    for s in ["Pocket", "Hat", "Block", "Tag", "Cite", "Verbatim"]:
        _add_style(doc, s)

    def p(style, text):
        para = doc.add_paragraph()
        para.style = doc.styles[style]
        para.add_run(text)

    p("Pocket", "Deterrence Advantage")
    p("Hat", "Nuclear War")
    p("Block", "Deterrence Works")

    p("Tag", "Deterrence prevents great power conflict")
    p("Cite", "Smith 2023, Professor of Political Science at MIT, Foreign Policy Journal, p. 42")
    p("Verbatim", "Nuclear deterrence has maintained strategic stability since the Cold War. "
                  "States with nuclear weapons have not fought direct wars against each other.")

    p("Tag", "Credibility is key to deterrence stability")
    p("Cite", "Jones 2022, Senior Fellow at RAND Corporation, Journal of Strategic Studies")
    p("Verbatim", "The credibility of the nuclear arsenal depends on political will as much as "
                  "capability. Adversaries calculate resolve as well as capability.")

    doc.save(str(path))


def make_generic(path: Path):
    """Generic debate doc: bold tags, italic cites, plain body text, 2 cards."""
    doc = Document()

    def bold(text):
        p = doc.add_paragraph()
        p.add_run(text).bold = True

    def italic(text):
        p = doc.add_paragraph()
        p.add_run(text).italic = True

    bold("Warming causes extinction")
    italic("Hansen 2023, Director of NASA GISS, Climate Science Review, vol. 45")
    doc.add_paragraph(
        "Climate change poses existential risks to human civilization. "
        "Runaway warming beyond 4 degrees Celsius would trigger cascading ecosystem collapse."
    )

    bold("Feedback loops accelerate collapse")
    italic("Mann 2021, Penn State University, Nature Climate Change journal")
    doc.add_paragraph(
        "Arctic methane feedback loops could accelerate warming beyond model predictions. "
        "The permafrost thaw releases stored carbon into the atmosphere."
    )

    doc.save(str(path))


def make_malformed(path: Path):
    """Edge-case doc: tag without cite, tag without body, empty paragraph."""
    doc = Document()
    for s in ["Tag", "Cite", "Verbatim"]:
        _add_style(doc, s)

    def p(style, text):
        para = doc.add_paragraph()
        para.style = doc.styles[style]
        para.add_run(text)

    # Card 1: tag with card text but NO cite
    p("Tag", "Tag without cite")
    p("Verbatim", "Card text with no cite above it.")

    # Card 2: tag with no body (immediately followed by next tag)
    p("Tag", "Tag without body")

    # Card 3: well-formed card right after the bodyless tag
    p("Tag", "Another tag right after")
    p("Cite", "Smith 2023, Some University")
    p("Verbatim", "This card has a cite and body but the previous tag had none.")

    # Empty paragraph — should be skipped
    doc.add_paragraph()

    doc.save(str(path))


if __name__ == "__main__":
    FIXTURES.mkdir(parents=True, exist_ok=True)
    make_verbatim(FIXTURES / "verbatim_sample.docx")
    make_generic(FIXTURES / "generic_sample.docx")
    make_malformed(FIXTURES / "malformed_sample.docx")
    print("Fixtures written to", FIXTURES)
