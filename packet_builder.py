"""Build a .docx card packet from selected DebateCard objects.

Preserves Verbatim-compatible paragraph styles where possible so the
output can be opened in Verbatim/CardMirror. Falls back to explicit
character formatting (bold/underline) when Verbatim styles are unavailable.
"""

from __future__ import annotations

import time
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from models import DebateCard

_DEFAULT_PACKETS_DIR = Path.home() / ".opencaselist-mcp" / "packets"


def _add_run(para, text: str, bold=False, italic=False, underline=False,
             font_size: Optional[int] = None, color: Optional[RGBColor] = None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    return run


def _set_para_style(para, style_name: str, doc: Document):
    """Apply a named style if it exists in the document, else skip."""
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        pass


def _add_heading(doc: Document, text: str, level: int = 2):
    h = doc.add_heading(text, level=level)
    return h


def _add_pocket_label(doc: Document, text: str):
    para = doc.add_paragraph()
    try:
        para.style = doc.styles["Pocket"]
    except KeyError:
        _add_run(para, text, bold=True, font_size=14)
        return para
    _add_run(para, text)
    return para


def _add_block_label(doc: Document, text: str):
    para = doc.add_paragraph()
    try:
        para.style = doc.styles["Block"]
    except KeyError:
        _add_run(para, text, bold=True, font_size=12)
        return para
    _add_run(para, text)
    return para


def _add_tag(doc: Document, text: str):
    para = doc.add_paragraph()
    try:
        para.style = doc.styles["Tag"]
    except KeyError:
        _add_run(para, text, bold=True, underline=True)
        return para
    _add_run(para, text)
    return para


def _add_cite(doc: Document, text: str):
    para = doc.add_paragraph()
    try:
        para.style = doc.styles["Cite"]
    except KeyError:
        _add_run(para, text, italic=True)
        return para
    _add_run(para, text)
    return para


def _add_card_text(doc: Document, text: str):
    para = doc.add_paragraph()
    try:
        para.style = doc.styles["Verbatim"]
    except KeyError:
        try:
            para.style = doc.styles["Normal"]
        except KeyError:
            pass
        _add_run(para, text)
        return para
    _add_run(para, text)
    return para


def _add_source_metadata(doc: Document, card: DebateCard):
    """Add a small metadata line after a card noting its origin."""
    parts = []
    if card.team:
        parts.append(card.team)
    if card.tournament:
        parts.append(card.tournament)
    if card.round:
        parts.append(f"R{card.round}")
    if card.side and card.side != "unknown":
        parts.append(card.side.upper())
    if card.source_file:
        parts.append(f"[{card.source_file}]")
    if card.wiki_url:
        parts.append(card.wiki_url)

    if parts:
        para = doc.add_paragraph()
        _add_run(para, "Source: " + " | ".join(parts), italic=True, font_size=8)


def build_card_packet(
    cards: List[DebateCard],
    output_path: Optional[Path] = None,
    group_by: str = "pocket",  # "pocket", "block", "source_file", "none"
    include_source_metadata: bool = True,
    title: str = "Evidence Packet",
) -> Dict:
    """
    Build a .docx card packet from a list of DebateCard objects.

    Returns {"success": True, "path": str, "card_count": int} or an error dict.
    """
    if not cards:
        return {"success": False, "error": "No cards provided."}

    if output_path is None:
        _DEFAULT_PACKETS_DIR.mkdir(parents=True, exist_ok=True)
        ts = int(time.time())
        output_path = _DEFAULT_PACKETS_DIR / f"packet_{ts}.docx"

    doc = Document()

    # Title page
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Generated {time.strftime('%Y-%m-%d %H:%M')} | {len(cards)} cards")
    doc.add_page_break()

    # Group cards
    groups: Dict[str, List[DebateCard]] = {}
    if group_by == "pocket":
        for card in cards:
            key = card.pocket or card.block or "Ungrouped"
            groups.setdefault(key, []).append(card)
    elif group_by == "block":
        for card in cards:
            key = card.block or card.pocket or "Ungrouped"
            groups.setdefault(key, []).append(card)
    elif group_by == "source_file":
        for card in cards:
            key = card.source_file or "Unknown Source"
            groups.setdefault(key, []).append(card)
    else:
        groups[""] = cards

    for group_label, group_cards in groups.items():
        if group_label:
            _add_block_label(doc, group_label)

        for card in group_cards:
            if card.tag:
                _add_tag(doc, card.tag)
            if card.cite:
                _add_cite(doc, card.cite)
            if card.card_text:
                _add_card_text(doc, card.card_text)
            if include_source_metadata:
                _add_source_metadata(doc, card)
            doc.add_paragraph()  # spacing

    try:
        doc.save(str(output_path))
        return {
            "success": True,
            "path": str(output_path),
            "card_count": len(cards),
            "group_count": len(groups),
        }
    except Exception as e:
        return {"success": False, "error": str(e)}
